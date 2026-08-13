from __future__ import annotations

import json
import sys
from collections import Counter
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


BODY_STYLE_NAMES = {
    "Normal",
    "Body Text",
    "正文",
    "正文文本",
    "Body Text 2",
    "Body Text 3",
    "List Paragraph",
    "列表段落",
}

EXCLUDED_STYLE_MARKERS = (
    "Title", "标题", "Heading", "题注", "Caption", "TOC", "目录",
    "Quote", "引用", "Bibliography", "参考文献", "Equation", "公式",
    "Code", "代码", "Footnote", "脚注", "Endnote", "尾注",
)


def iter_paragraphs(parent):
    for paragraph in parent.paragraphs:
        yield paragraph
    for table in parent.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from iter_paragraphs(cell)


def iter_body_paragraphs(doc: Document):
    """Yield thesis content from the first level-1 heading, plus table text."""
    started = False
    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name if paragraph.style is not None else "Normal"
        if not started and style_name == "Heading 1":
            started = True
        if started:
            yield paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from iter_paragraphs(cell)


def style_is_body(paragraph) -> bool:
    style = paragraph.style
    name = style.name if style is not None else "Normal"
    if any(marker.lower() in name.lower() for marker in EXCLUDED_STYLE_MARKERS):
        return False
    return name in BODY_STYLE_NAMES or name.startswith("Normal") or name.startswith("正文")


def color_value(run):
    color = run.font.color.rgb
    return str(color) if color is not None else None


def audit(doc: Document):
    stats = Counter()
    styles = Counter()
    examples = []
    for paragraph in iter_body_paragraphs(doc):
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = paragraph.style.name if paragraph.style is not None else "Normal"
        styles[style_name] += 1
        if not style_is_body(paragraph):
            continue
        for run in paragraph.runs:
            if not run.text.strip():
                continue
            key = (
                run.font.name,
                run._element.rPr.rFonts.get(qn("w:eastAsia")) if run._element.rPr is not None and run._element.rPr.rFonts is not None else None,
                run.font.size.pt if run.font.size else None,
                color_value(run),
                run.bold,
                run.italic,
                run.font.superscript,
                run.font.subscript,
            )
            stats[str(key)] += len(run.text)
            if len(examples) < 15:
                examples.append({"style": style_name, "text": run.text[:80], "format": key})
    return {"styles": styles.most_common(), "body_run_formats_by_characters": stats.most_common(), "examples": examples}


def set_run_font(run, font_name: str, size_pt: float, color: RGBColor):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), font_name)
    # Remove theme colors/tints so explicit black wins in all Word renderers.
    color_el = rpr.find(qn("w:color"))
    if color_el is not None:
        for attr in ("themeColor", "themeTint", "themeShade"):
            color_el.attrib.pop(qn(f"w:{attr}"), None)


def normalize(doc: Document):
    # Chinese academic-body convention: 宋体、小四(12 pt)、黑色。
    font_name = "宋体"
    size_pt = 12
    black = RGBColor(0, 0, 0)

    for style in doc.styles:
        if style.type != WD_STYLE_TYPE.PARAGRAPH:
            continue
        if style.name in BODY_STYLE_NAMES or style.name.startswith("Normal") or style.name.startswith("正文"):
            if not any(marker.lower() in style.name.lower() for marker in EXCLUDED_STYLE_MARKERS):
                style.font.name = font_name
                style.font.size = Pt(size_pt)
                style.font.color.rgb = black
                rpr = style.element.get_or_add_rPr()
                rfonts = rpr.rFonts
                if rfonts is None:
                    rfonts = OxmlElement("w:rFonts")
                    rpr.insert(0, rfonts)
                for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
                    rfonts.set(qn(f"w:{attr}"), font_name)

    changed_runs = 0
    changed_paragraphs = 0
    for paragraph in iter_body_paragraphs(doc):
        if not paragraph.text.strip() or not style_is_body(paragraph):
            continue
        changed_paragraphs += 1
        for run in paragraph.runs:
            if not run.text.strip():
                continue
            # Preserve semantic emphasis and scientific typography while making
            # the requested font/size/color uniform.
            set_run_font(run, font_name, size_pt, black)
            changed_runs += 1
    return {"font": font_name, "size_pt": size_pt, "color": "000000", "paragraphs": changed_paragraphs, "runs": changed_runs}


def main():
    if len(sys.argv) != 4:
        raise SystemExit("usage: audit_and_normalize_thesis_docx.py INPUT OUTPUT REPORT")
    source, output, report_path = map(Path, sys.argv[1:])
    doc = Document(source)
    before = audit(doc)
    changes = normalize(doc)
    doc.save(output)
    verified = Document(output)
    after = audit(verified)
    report_path.write_text(json.dumps({"source": str(source), "output": str(output), "changes": changes, "before": before, "after": after}, ensure_ascii=False, indent=2), encoding="utf-8")


if __name__ == "__main__":
    main()
